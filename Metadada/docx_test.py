from docx import Document
import arcpy
# from arcpy import metadata as md
import os
import xml.etree.ElementTree as ET
from docx.shared import Pt


variaveis = {

    "Prec": "Precipitação",
    "OVrn": "Ocorrência de veranicos",
    "PChv": "Duração do período chuvoso",
    "DChv": "Dias chuvosos",
    "ISec": "Índice de seca",
    "TMax": "Temperatura máxima",
    "TMin": "Temperatura mínima",
    "TMed": "Temperatura média",
    "IHid": "Índice hídrico",
    "Umid": "Umidade relativa do ar",
    "Insl": "Insolação",
    "Nebl": "Nebulosidade",
    "RdSG": "Radiação solar global",
    "Evap": "Evaporação",
    "EtoR": "Evapotranspiração de referência",
    "EHid": "Excedente hídrico",
    "DHid": "Deficiência hídrica",
    "RClm": "Regionalização climática"

}


document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

path_gdb = r""

path_gdb_250 = r""

path_base_1M = r""

path_base_250 = r""

path_base_relatorio = r""

featureclasses = [

    # os.path.join(path_gdb, "Deficiencia_hidrica_GCS", "DHidAnual_1M_GCS_a"),
    # os.path.join(path_gdb, "Evaporacao_GCS", "EvapMM01_1M_GCS_a"),
    # os.path.join(path_gdb, "Evapotranspiracao_GCS", "EtoRMM01_1M_GCS_a"),
    # os.path.join(path_gdb, "Excedente_hidrico_GCS", "EHidAnual_1M_GCS_a"),
    # os.path.join(path_gdb, "Indice_Hidrico_GCS", "IHid_1M_GCS_a"),
    # os.path.join(path_gdb, "Indice_Seca_GCS", "ISec_1M_GCS_a"),
    # os.path.join(path_gdb, "Insolacao_GCS", "InslMM01_1M_GCS_a"),
    # os.path.join(path_gdb, "Precipitacao_GCS", "PChv1990a1991_1M_GCS_a"),
    # os.path.join(path_gdb, "Precipitacao_GCS", "PrecMM01_1M_GCS_a"),
    # os.path.join(path_gdb, "Nebulosidade_GCS", "Nebl1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Precipitacao_GCS", "DChv1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Radiacao_Solar_Global_GCS", "RdSG2007_1M_GCS_a"),
    # os.path.join(path_gdb, "Regionalizacao_climatica_GCS", "RClm_1M_GCS_a"),
    # os.path.join(path_gdb, "Temperatura_GCS", "TMax1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Temperatura_GCS", "TMed1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Temperatura_GCS", "TMin1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Umidade_GCS", "Umid1990_1M_GCS_a"),
    # os.path.join(path_gdb, "Veranicos_GCS", "OVrn1990a1991_1M_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "DChv1990_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "EtoRMM01_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "IHid_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "ISec_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "OVrn1990a1991_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "PChv1990a1991_250_GCS_a"),
    # os.path.join(path_gdb_250, "TO_250", "PrecMM01_250_GCS_a"),
    # os.path.join(path_base_1M, "Area_Indigena"),
    # os.path.join(path_base_1M, "Ferrovias"),
    # os.path.join(path_base_1M, "Hidreletrica"),
    # os.path.join(path_base_1M, "Hidrografia"),
    # os.path.join(path_base_1M, "Ilhas"),
    # os.path.join(path_base_1M, "Lagos_UHE"),
    # os.path.join(path_base_1M, "Limite_Estado"),
    # os.path.join(path_base_1M, "Limite_Municipal_Linha"),
    # os.path.join(path_base_1M, "Localidades"),
    # os.path.join(path_base_1M, "Massas_Agua"),
    # os.path.join(path_base_1M, "Rodovias"),
    # os.path.join(path_base_1M, "Sedes_Municipais"),
    # os.path.join(path_base_1M, "Unidades_Conservacao"),
    # os.path.join(path_base_250, "Bases", "Articulacao_250mi"),
    # os.path.join(path_base_250, "F172_GCS", "Hidrografia_Linha_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Limite_Municipal_Linha_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Localidades_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Massa_Agua_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Rodovias_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Sedes_Municipais_172_GCS"),
    # os.path.join(path_base_250, "F172_GCS", "Unidades_Conservacao_172_GCS"),
    # os.path.join(path_base_250, "F173_GCS", "Area_Indigena_173_GCS"),
    # os.path.join(path_base_250, "F199_GCS", "Hidreletrica_199_GCS"),
    # os.path.join(path_base_250, "F200_GCS", "Ferrovias_200_GCS"),
    # os.path.join(path_base_250, "F226_GCS", "Lagos_UHE_226_GCS"),
    # os.path.join(path_base_relatorio, "Agro_IBGE_2006"),
    # os.path.join(path_base_relatorio, "Agro_IBGE_2017"),
    # os.path.join(path_base_relatorio, "CAD"),
    # os.path.join(path_base_relatorio, "Estacoes_Meteorologicas"),
    # os.path.join(path_base_relatorio, "Regionalizacao_Koppen_Geiger"),
    ]

# arcpy.env.workspace = gdb
# featureclasses = arcpy.ListFeatureClasses()
i = 1
for fc in featureclasses:
    desc = arcpy.Describe(fc)
    dataset_fc = desc.path.split("\\")[-1]

    out_shp = r""
    arcpy.env.workspace = path_gdb
    arcpy.FeatureClassToShapefile_conversion(fc, out_shp)
    shp = os.path.join(out_shp, desc.name + ".shp.xml")
    print(shp)

    tree_1 = ET.parse(shp)
    root = tree_1.getroot()

    p0 = document.add_paragraph()
    p0.add_run(str(i)).bold = True
    p0.add_run(". ").bold = True
    p0.add_run(fc).bold = True

    p1 = document.add_paragraph()
    p1.add_run("Feature Dataset: ").bold = True
    p1.add_run(dataset_fc)

    p2 = document.add_paragraph()
    p2.add_run("Feature Classes: ").bold = True
    p2.add_run(desc.name)
    # listar todas as feature classes deste dado?

    p2 = document.add_paragraph()
    p2.add_run("IDENTIFICAÇÃO").underline = True

    p2 = document.add_paragraph()
    p2.add_run("Título: ").bold = True
    p2.add_run(root.findall("./dataIdInfo/idCitation/resTitle")[0].text)

    # p3 = document.add_paragraph()
    # p3.add_run("Data: ").bold = True
    # p3.add_run(root.findall("./dataIdInfo/idCitation/date/createDate")[0].text)
    # p3.add_run(" criação")

    # p3 = document.add_paragraph()
    # p3.add_run("Série: ").bold = True
    # p3.add_run(root.findall("./dataIdInfo/idCitation/datasetSeries/seriesName")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("ISBN: ").bold = True
    p3.add_run("Não Identificado")

    p3 = document.add_paragraph()
    p3.add_run("Resumo: ").bold = True
    p3.add_run(root.findall("./dataIdInfo/idPurp")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Objetivo: ").bold = True
    p3.add_run(root.findall("./dataIdInfo/idAbs")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Créditos: ").bold = True
    p3.add_run(root.findall("./dataIdInfo/idCredit")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Status: ").bold = True
    p3.add_run("Não Identificado")

    p3 = document.add_paragraph()
    p3.add_run("Responsável: ").bold = True
    p3.add_run(root.findall("./dataIdInfo/idCredit")[0].text)
    # p3.add_run(", ")
    # p3.add_run("produtor").bold = True
    # p3.add_run(".")

    p3 = document.add_paragraph()
    p3.add_run("Palavras Chave: ").bold = True
    p3.add_run(root.findall("./dataIdInfo/searchKeys/keyword")[0].text)

    p2 = document.add_paragraph()
    p2.add_run("IDENTIFICAÇÃO DO CDG").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Categoria Temática: ").bold = True
    p3.add_run("Planejamento e Cadastro")

    p3 = document.add_paragraph()
    p3.add_run("Idioma: ").bold = True
    p3.add_run("Português (Brasil)")

    p3 = document.add_paragraph()
    p3.add_run("Norma de Codificação de Caracteres: ").bold = True
    p3.add_run("Uft8")

    p3 = document.add_paragraph()
    p3.add_run("Escala: ").bold = True
    p3.add_run("1:1.000.000")

    p3 = document.add_paragraph()
    p3.add_run("Tipo de Representação Espacial: ").bold = True
    p3.add_run("Vetorial")

    p3 = document.add_paragraph()
    p3.add_run("Extensão Geográfica: ").bold = True
    p3.add_run("Norte: " + root.findall("./Esri/DataProperties/itemProps/nativeExtBox/northBL")[0].text + " ")
    p3.add_run("Sul: " + root.findall("./Esri/DataProperties/itemProps/nativeExtBox/southBL")[0].text + " ")
    p3.add_run("Leste: " + root.findall("./Esri/DataProperties/itemProps/nativeExtBox/eastBL")[0].text + " ")
    p3.add_run("Oeste: " + root.findall("./Esri/DataProperties/itemProps/nativeExtBox/westBL")[0].text)

    p2 = document.add_paragraph()
    p2.add_run("RESTRIÇÕES").underline = True

    # p3 = document.add_paragraph()
    # p3.add_run("Restrições Legais: ").bold = True
    # p3.add_run(root.findall("./dataIdInfo/resConst/Consts/useLimit")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Restrições de Segurança: ").bold = True
    p3.add_run("Não Classificado")

    p2 = document.add_paragraph()
    p2.add_run("QUALIDADE").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Linhagem: ").bold = True
    p3.add_run(root.findall("./dqInfo/dataLineage/statement")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Fonte do dados / Descrição da Fonte: ").bold = True
    p3.add_run(root.findall("./dqInfo/dataLineage/dataSource/srcDesc")[0].text)

    # p3 = document.add_paragraph()
    # p3.add_run("Etapas do Processo: ").bold = True
    # p3.add_run(root.findall("./dqInfo/dataLineage/prcStep/stepDesc")[0].text)

    p2 = document.add_paragraph()
    p2.add_run("MANUTENÇÃO").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Frequência de Manutenção e Atualização: ").bold = True
    p3.add_run("Conforme necessidade")

    p2 = document.add_paragraph()
    p2.add_run("REPRESENTAÇÃO ESPACIAL VETORIAL").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Representação Espacial Vetorial: ").bold = True
    p3.add_run("Poligono")

    p2 = document.add_paragraph()
    p2.add_run("SISTEMA DE REFERÊNCIA").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Sistema de Referência: ").bold = True
    p3.add_run("GCS SIRGAS 2000")

    p2 = document.add_paragraph()
    p2.add_run("DISTRIBUIÇÃO").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Nome da organização: ").bold = True
    p3.add_run("Secretaria da Fazenda...")
    p3.add_run("Função: ").bold = True
    p3.add_run("Produtor")
    p3.add_run("Telefones:")
    p3.add_run("Endereço:")
    p3.add_run("CEP:")
    p3.add_run("Cidade:")
    p3.add_run("UF:")
    p3.add_run("Pais: Brasil")
    p3.add_run("e-mails: ")

    p3 = document.add_paragraph()
    p3.add_run("Nome da organização: ").bold = True
    p3.add_run("Secretaria da Fazenda...")
    p3.add_run("Função: ").bold = True
    p3.add_run("Produtor")
    p3.add_run("Telefones:")
    p3.add_run("Endereço:")
    p3.add_run("CEP:")
    p3.add_run("Cidade:")
    p3.add_run("UF:")
    p3.add_run("Pais: Brasil")
    p3.add_run("e-mails: ")

    p2 = document.add_paragraph()
    p2.add_run("METAMETADADOS").underline = True

    p3 = document.add_paragraph()
    p3.add_run("Data de criação dos metadados: ").bold = True
    p3.add_run(root.findall("./Esri/CreaDate")[0].text)

    p3 = document.add_paragraph()
    p3.add_run("Idioma: ").bold = True
    p3.add_run("Português (Brasil)")

    p3 = document.add_paragraph()
    p3.add_run("Norma de Codificação de Caracteres: ").bold = True
    p3.add_run("Uft8")

    p3 = document.add_paragraph()
    p3.add_run("Nível Hierárquico: ").bold = True
    p3.add_run("Conjunto de Dados Geograficos")

    p3 = document.add_paragraph()
    p3.add_run("Identificador: ").bold = True
    p3.add_run(root.find("mdFileID").text)

    p3 = document.add_paragraph()
    p3.add_run("Nome da organização: ").bold = True
    p3.add_run("Secretaria da Fazenda...")
    p3.add_run("Função: ").bold = True
    p3.add_run("Produtor")
    p3.add_run("Telefones:")
    p3.add_run("Endereço:")
    p3.add_run("CEP:")
    p3.add_run("Cidade:")
    p3.add_run("UF:")
    p3.add_run("Pais: Brasil")
    p3.add_run("e-mails: ")

    p3 = document.add_paragraph()
    p3.add_run("Designação da Norma e perfil de metadados: ").bold = True
    p3.add_run("ISO 19115/CONCAR")

    p3 = document.add_paragraph()
    p3.add_run("Versão da norma de metadados: ").bold = True
    p3.add_run("2009")

    document.add_page_break()

    # arcpy.Delete_management(shp)
    i = i + 1

document.add_page_break()

document.save(r'')

